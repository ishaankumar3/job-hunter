#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
====================================================
 CV & Cover Letter Tailoring Engine — Ishaan Kumar
====================================================
 For each job:
  1. Sends job description to Claude API
  2. Gets back tailored CV content + cover letter
  3. Generates professional UK-format Word documents
  4. Returns file paths for email attachment
====================================================
"""

import os, json, re, logging
from datetime import datetime

log = logging.getLogger(__name__)

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")

def _check_dependencies():
    """Test all required packages and log clearly what is/isn't available."""
    ok = True
    try:
        import anthropic
        log.info("[DEP] anthropic: OK")
    except ImportError as e:
        log.error(f"[DEP] anthropic MISSING: {e} -- run: pip install anthropic")
        ok = False
    try:
        from docx import Document
        log.info("[DEP] python-docx: OK")
    except ImportError as e:
        log.error(f"[DEP] python-docx MISSING: {e} -- run: pip install python-docx")
        ok = False
    return ok

# ═══════════════════════════════════════════════════════════════════
#  ISHAAN'S MASTER CV DATA — extracted from Master_CV V9
# ═══════════════════════════════════════════════════════════════════

MASTER_CV = {
    "name": "ISHAAN KUMAR",
    "contact": {
        "location": "London, United Kingdom",
        "phone": "+44 7586 603561",
        "email": "ishaankumar3@gmail.com",
        "linkedin": "linkedin.com/in/ishaankumar3",
    },
    "profile": """Highly driven Water / Mechanical Design Engineer with 3.5+ years of professional experience, 
including 2+ years in the UK water and wastewater infrastructure sector. Experienced across concept, 
detailed design, and construction support, with strong capability in CAD/BIM delivery, 
multidisciplinary coordination, and regulatory compliance. Background includes work on UK AMP7/AMP8 
projects, water networks, wastewater systems, pumping stations, and infrastructure upgrades. 
Strong working knowledge of NAV sites, Section 38/104/106/185/278, and incumbent water company 
specifications (Thames Water, Affinity Water, South East Water, Wessex Water). Currently delivering 
end-to-end NAV and Self-Lay schemes with full technical approval responsibility for projects 
valued at over £20M.""",

    "key_achievements": [
        "Delivered over 100 km of compliant water mains design across multiple UK developments",
        "Successfully secured technical approvals from Thames Water and Affinity Water on first submission for numerous schemes",
        "Managed a concurrent project portfolio exceeding £20M, delivering designs within programme and budget",
        "Designed and delivered 100+ km of water distribution networks across 20+ residential and mixed-use development sites",
        "Achieved consistent cost savings through value engineering principles across multiple regulated water infrastructure projects",
        "Developed standardised design workflows adopted across the business, improving team efficiency",
        "Mentored junior engineers and standardised CAD templates and documentation formats",
    ],

    "technical_skills": {
        "engineering_design": [
            "Water Network Design (NAV & Self-Lay end-to-end)",
            "Hydraulic Design & Network Modelling (WaterGEMS)",
            "Developer Services — Section 104/185/38/278 applications",
            "Pumping Station concept & detailed design support",
            "Mechanical Systems Design (pipework & equipment layouts)",
            "Design Risk Assessments (DRA)",
            "Constructability & Buildability Reviews",
            "Value Engineering",
            "CDM 2015 Design Stage",
            "SuDS & Drainage Design",
        ],
        "software": [
            "WaterGEMS (hydraulic modelling — practical project experience)",
            "AutoCAD (Advanced — 2D & 3D, certified)",
            "Revit (Intermediate — BIM coordination, families, views)",
            "Civil 3D (foundation level)",
            "ArcGIS Pro (spatial analysis & infrastructure mapping)",
            "InfraWorks (drainage & infrastructure planning)",
            "Navisworks (clash detection)",
            "SolidWorks / CREO / CATIA",
            "Microsoft Excel, Power BI",
        ],
        "standards": [
            "UK Water Industry Standards & Water UK design standards",
            "NAV & Self-Lay (SLP) end-to-end process management",
            "Thames Water & Affinity Water technical approval processes",
            "Sewers for Adoption (SFA) & WRAS compliance",
            "CDM 2015, DWI & WRs guidance, Ofwat",
            "British Standards (BS) & ISO quality processes",
            "Section 38, 104, 185, 278 agreements",
        ],
    },

    "experience": [
        {
            "title": "Water Design Engineer",
            "company": "UKPS",
            "location": "London Area, United Kingdom",
            "dates": "Oct 2024 – Present",
            "bullets": [
                "Independently deliver end-to-end design responsibility for NAV and Self-Lay water infrastructure schemes from feasibility through to technical approval and construction support",
                "Designed and delivered 100+ km of water distribution networks across 20+ residential and mixed-use development sites",
                "Managed projects with a combined portfolio value exceeding £20 million across 7–20 concurrent projects",
                "Act as the primary technical interface with adopting authorities including Thames Water, Affinity Water, Icosa and other incumbent companies, securing technical approvals on first submission",
                "Prepare and submit full technical approval packages including design calculations, drawings, cost estimates, risk assessments and compliance documentation",
                "Ensure designs achieve first-time compliance with SFA, Water UK, WRAS and company-specific technical standards, minimising approval iterations",
                "Lead utility and multidisciplinary coordination, resolving clashes with electrical, gas, highways, drainage, landscaping and MEP infrastructure to deliver buildable solutions",
                "Conduct site investigations to assess existing water infrastructure prior to network modification or disconnection",
                "Produce formal design risk assessments integrating safety, environmental, and constructability considerations",
                "Applied value engineering principles achieving consistent cost savings through material selection, routing optimisation, and constructability-led decisions",
                "Mentored junior engineers and collaborated on standardising CAD templates and documentation formats",
                "Prepare technical design drawings, layouts, sections, and profiles using AutoCAD and Revit",
                "Support hydraulic and network design activities using WaterGEMS for pressure, flow, and fire flow analysis",
            ],
        },
        {
            "title": "Assistant Water Design Engineer",
            "company": "UKPS",
            "location": "Greater London, United Kingdom",
            "dates": "Oct 2023 – Oct 2024",
            "bullets": [
                "Supported senior engineers in delivery of water network designs for NAV and Self-Lay schemes",
                "Produced CAD drawings, design documents, and technical specifications in line with client and regulatory standards",
                "Assisted with design risk assessments, project coordination, and document control",
                "Managed multiple work packages simultaneously, demonstrating strong time management and prioritisation",
                "Gained hands-on experience across AutoCAD, WaterGEMS, and engineering documentation",
                "Gained extensive knowledge of NAVs (New Appointments and Variations) and incumbent water supply operations",
            ],
        },
        {
            "title": "Design Engineer",
            "company": "Bellis Hardware Private Limited",
            "location": "Noida, India",
            "dates": "Sep 2021 – Jul 2022",
            "bullets": [
                "Designed architectural hardware components for global clients including Häfele Dubai and 5/7-star hotel projects worldwide",
                "Led the design department overseeing concept-to-production design activities",
                "Created detailed 3D models, prototypes, and 2D manufacturing drawings using CREO and AutoCAD",
                "Developed Bills of Materials (BOMs) and supported cost estimation and manufacturing process selection",
            ],
        },
    ],

    "education": [
        {
            "degree": "MSc Advanced Mechanical Engineering",
            "institution": "Cranfield University, United Kingdom",
            "details": "Modules: Risk & Reliability Engineering, Structural Integrity, Fluid Mechanics, Computational Fluid Dynamics, Engineering Project Management. Thesis: Assessing the Performance of Floating Solar Panels in Variable Wave Conditions",
        },
        {
            "degree": "BTech Mechanical Engineering",
            "institution": "Manipal University Jaipur, India",
            "details": "Core mechanical engineering, design, manufacturing and materials. Thesis: Development and Testing of Eutectic Mixture for Thermal Energy Storage",
        },
    ],

    "certifications": [
        "ProQual Level 5 Award — Understanding Developer Services for Water & Environmental Industries (Sep 2024)",
        "EUSR — Clean Water & Waste Water (UK Water Industry Safety Certification)",
        "AutoCAD 2D/3D Certification — CADD Centre (Aug 2022)",
        "Member — Institution of Engineering and Technology (MIET)",
        "Actively working towards Chartered Engineer (CEng) status",
        "ArcGIS Pro Essential Training (2026)",
        "Learning Revit 2026 — LinkedIn Learning (Jan 2026)",
        "Navisworks Essential Training — LinkedIn Learning (Jan 2026)",
        "InfraWorks 2020 & Drainage Design — LinkedIn Learning (2026)",
        "BIM: Designing Sustainable HVAC Systems with Revit — LinkedIn Learning (Jan 2026)",
        "Power BI Essential Training (2026)",
        "IELTS — C1 Proficiency (Band 7)",
        "Full UK driving licence",
    ],

    "languages": "German (conversational), Spanish (conversational)",
    "right_to_work": "Right to work in the UK (company-sponsored visa). Open to UK and international opportunities.",
}


# ═══════════════════════════════════════════════════════════════════
#  STEP 1 — CALL CLAUDE API TO TAILOR CV + COVER LETTER
# ═══════════════════════════════════════════════════════════════════

def get_tailored_content(job: dict) -> dict | None:
    """Call Claude API to analyse job and return tailored CV + cover letter content."""
    if not ANTHROPIC_API_KEY:
        log.warning("ANTHROPIC_API_KEY not set — skipping CV tailoring")
        return None

    try:
        import anthropic as _anthropic
    except ImportError:
        log.error("[FAIL] anthropic not installed — pip install anthropic==0.30.0")
        return None
    client = _anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    prompt = f"""You are an expert UK CV writer and career coach specialising in water engineering roles.

I need you to tailor a CV and write a cover letter for the following job application.

## JOB DETAILS
Title: {job['title']}
Company: {job['company']}
Location: {job['location']}
Salary: {job['salary']}
Job Description:
{job['description']}

## CANDIDATE'S MASTER CV DATA
{json.dumps(MASTER_CV, indent=2)}

## YOUR TASK

Respond with ONLY a valid JSON object (no markdown, no explanation, just JSON) with this exact structure:

{{
  "cv": {{
    "profile": "2-3 sentence tailored professional profile that directly mirrors the job description language and requirements. Start with job title from the advert. Include key matching skills.",
    "top_skills": ["skill 1", "skill 2", "skill 3", "skill 4", "skill 5", "skill 6", "skill 7", "skill 8"],
    "experience_bullets": {{
      "Water Design Engineer": [
        "tailored bullet 1 using keywords from job description",
        "tailored bullet 2",
        "tailored bullet 3",
        "tailored bullet 4",
        "tailored bullet 5",
        "tailored bullet 6"
      ],
      "Assistant Water Design Engineer": [
        "tailored bullet 1",
        "tailored bullet 2",
        "tailored bullet 3"
      ]
    }},
    "relevant_certifications": ["cert 1", "cert 2", "cert 3", "cert 4"],
    "ats_keywords": ["keyword1", "keyword2", "keyword3"]
  }},
  "cover_letter": {{
    "opening_paragraph": "Strong opening paragraph that mentions the exact role and company. Show genuine interest in THIS company specifically. 3-4 sentences.",
    "why_me_paragraph": "Paragraph explaining why Ishaan is perfect for this role with specific evidence. Reference £20M portfolio, first-time approvals, 100km of water mains etc. 3-4 sentences.",
    "bullet_skills": [
      {{
        "skill": "Skill name matching job requirement",
        "example": "Specific real example from CV with quantified result"
      }},
      {{
        "skill": "Skill name matching job requirement",
        "example": "Specific real example from CV with quantified result"
      }},
      {{
        "skill": "Skill name matching job requirement",
        "example": "Specific real example from CV with quantified result"
      }},
      {{
        "skill": "Skill name matching job requirement",
        "example": "Specific real example from CV with quantified result"
      }}
    ],
    "closing_paragraph": "Enthusiastic closing paragraph. Express eagerness to discuss further. Professional but warm. 2-3 sentences.",
    "company_research_note": "1-2 sentence note about the company or sector based on what you know, to make the letter feel researched and personal."
  }}
}}

Rules:
- Use EXACT keywords from the job description throughout (for ATS)
- Keep bullets concise and action-led (start with strong verbs)
- Every claim must be backed by real evidence from the master CV
- Cover letter must feel human and personal — NOT generic
- UK English spelling throughout
- Do NOT invent experience that is not in the master CV"""

    try:
        message = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=3000,
            messages=[{"role": "user", "content": prompt}]
        )

        raw = message.content[0].text.strip()
        # Strip any accidental markdown fences
        raw = re.sub(r"^```json\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        data = json.loads(raw)
        log.info(f"[OK] Tailored content generated for: {job['title']} at {job['company']}")
        return data

    except json.JSONDecodeError as e:
        log.error(f"JSON parse error for {job['title']}: {e}")
        return None
    except Exception as e:
        log.error(f"API error for {job['title']}: {e}")
        return None


# ═══════════════════════════════════════════════════════════════════
#  STEP 2 — BUILD TAILORED CV (UK FORMAT, ATS-FRIENDLY)
# ═══════════════════════════════════════════════════════════════════

def _set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_section_heading(doc, text):
    """Add a blue section heading with underline rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    _set_font(run, size=10, bold=True, color=(0, 70, 127))
    # Bottom border as divider
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "00467F")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_bullet(doc, text, indent=0.4):
    """Add a bullet point paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _set_font(run, size=10)
    return p


def build_cv_docx(job: dict, tailored: dict, output_path: str):
    """Generate a professional UK-format tailored CV as a Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        log.error(f"[FAIL] python-docx not installed: {e} — pip install python-docx")
        raise
    doc = Document()

    # ── Page margins ───────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    cv_data = tailored["cv"]

    # ══════════════════════════════════════════════════════════
    #  HEADER — Name & Contact
    # ══════════════════════════════════════════════════════════
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(MASTER_CV["name"])
    _set_font(name_run, size=20, bold=True, color=(0, 70, 127))

    contact = MASTER_CV["contact"]
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(6)
    contact_run = contact_p.add_run(
        f"{contact['location']}  |  {contact['phone']}  |  {contact['email']}  |  {contact['linkedin']}"
    )
    _set_font(contact_run, size=9, color=(80, 80, 80))

    # Horizontal rule
    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(0)
    rule_p.paragraph_format.space_after = Pt(8)
    pPr = rule_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "00467F")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ══════════════════════════════════════════════════════════
    #  PROFESSIONAL PROFILE
    # ══════════════════════════════════════════════════════════
    _add_section_heading(doc, "Professional Profile")
    profile_p = doc.add_paragraph()
    profile_p.paragraph_format.space_after = Pt(4)
    profile_run = profile_p.add_run(cv_data["profile"])
    _set_font(profile_run, size=10)

    # ══════════════════════════════════════════════════════════
    #  CORE SKILLS
    # ══════════════════════════════════════════════════════════
    _add_section_heading(doc, "Core Technical Skills")

    # Display skills in 2 columns using a table
    skills = cv_data.get("top_skills", [])
    mid = (len(skills) + 1) // 2
    col1 = skills[:mid]
    col2 = skills[mid:]

    if skills:
        tbl = doc.add_table(rows=max(len(col1), len(col2)), cols=2)
        tbl.style = "Table Grid"
        for row in tbl.rows:
            for cell in row.cells:
                cell._tc.get_or_add_tcPr()
                # Remove borders
                tcBorders = OxmlElement("w:tcBorders")
                for border_name in ["top", "left", "bottom", "right"]:
                    border = OxmlElement(f"w:{border_name}")
                    border.set(qn("w:val"), "none")
                    tcBorders.append(border)
                cell._tc.tcPr.append(tcBorders)

        for i, skill in enumerate(col1):
            cell = tbl.cell(i, 0)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(">> " + skill)
            _set_font(run, size=10)

        for i, skill in enumerate(col2):
            if i < len(tbl.rows):
                cell = tbl.cell(i, 1)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(">> " + skill)
                _set_font(run, size=10)

    # ══════════════════════════════════════════════════════════
    #  PROFESSIONAL EXPERIENCE
    # ══════════════════════════════════════════════════════════
    _add_section_heading(doc, "Professional Experience")

    exp_bullets = cv_data.get("experience_bullets", {})

    for exp in MASTER_CV["experience"]:
        # Job title + company + dates line
        job_p = doc.add_paragraph()
        job_p.paragraph_format.space_before = Pt(6)
        job_p.paragraph_format.space_after = Pt(1)

        title_run = job_p.add_run(exp["title"])
        _set_font(title_run, size=11, bold=True, color=(0, 70, 127))

        sep_run = job_p.add_run("  |  ")
        _set_font(sep_run, size=10)

        company_run = job_p.add_run(exp["company"])
        _set_font(company_run, size=10, bold=True)

        # Location + dates on same line right-aligned via tab
        loc_p = doc.add_paragraph()
        loc_p.paragraph_format.space_after = Pt(3)
        loc_run = loc_p.add_run(f"{exp['location']}  |  {exp['dates']}")
        _set_font(loc_run, size=9, color=(100, 100, 100))
        loc_run.font.italic = True

        # Get tailored bullets for this role, or fall back to master CV bullets
        bullets = exp_bullets.get(exp["title"], exp["bullets"])
        for bullet in bullets[:8]:  # max 8 bullets per role
            _add_bullet(doc, bullet)

    # ══════════════════════════════════════════════════════════
    #  EDUCATION
    # ══════════════════════════════════════════════════════════
    _add_section_heading(doc, "Education")

    for edu in MASTER_CV["education"]:
        edu_p = doc.add_paragraph()
        edu_p.paragraph_format.space_before = Pt(4)
        edu_p.paragraph_format.space_after = Pt(1)
        deg_run = edu_p.add_run(edu["degree"])
        _set_font(deg_run, size=10, bold=True)

        inst_p = doc.add_paragraph()
        inst_p.paragraph_format.space_after = Pt(2)
        inst_run = inst_p.add_run(edu["institution"])
        _set_font(inst_run, size=10, color=(80, 80, 80))
        inst_run.font.italic = True

        detail_p = doc.add_paragraph()
        detail_p.paragraph_format.space_after = Pt(4)
        detail_run = detail_p.add_run(edu["details"])
        _set_font(detail_run, size=9, color=(100, 100, 100))

    # ══════════════════════════════════════════════════════════
    #  CERTIFICATIONS (relevant ones)
    # ══════════════════════════════════════════════════════════
    _add_section_heading(doc, "Certifications & Professional Membership")

    relevant_certs = cv_data.get("relevant_certifications", MASTER_CV["certifications"][:6])
    for cert in relevant_certs:
        _add_bullet(doc, cert)

    # ══════════════════════════════════════════════════════════
    #  ADDITIONAL
    # ══════════════════════════════════════════════════════════
    _add_section_heading(doc, "Additional Information")
    add_p = doc.add_paragraph()
    add_p.paragraph_format.space_after = Pt(2)
    add_run = add_p.add_run(
        f"{MASTER_CV['right_to_work']}  |  Languages: {MASTER_CV['languages']}"
    )
    _set_font(add_run, size=10)

    doc.save(output_path)
    log.info(f"[OK] CV saved: {output_path}")


# ═══════════════════════════════════════════════════════════════════
#  STEP 3 — BUILD COVER LETTER (UK FORMAT, PERSONAL & HUMAN)
# ═══════════════════════════════════════════════════════════════════

def build_cover_letter_docx(job: dict, tailored: dict, output_path: str):
    """Generate a professional UK cover letter as a Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        log.error(f"[FAIL] python-docx not installed: {e} — pip install python-docx")
        raise
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    cl_data = tailored["cover_letter"]
    contact = MASTER_CV["contact"]
    today = datetime.now().strftime("%d %B %Y")

    # ── Sender details (right-aligned) ──────────────────────
    for line in [MASTER_CV["name"], contact["phone"], contact["email"],
                 contact["linkedin"], contact["location"]]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        _set_font(run, size=10)

    doc.add_paragraph()  # spacer

    # ── Date ──────────────────────────────────────────────────
    date_p = doc.add_paragraph()
    date_p.paragraph_format.space_after = Pt(8)
    date_run = date_p.add_run(today)
    _set_font(date_run, size=10)

    # ── Recipient details ─────────────────────────────────────
    for line in ["Hiring Manager", job["company"], job["location"]]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        _set_font(run, size=10)

    doc.add_paragraph()  # spacer

    # ── Subject line ──────────────────────────────────────────
    subj_p = doc.add_paragraph()
    subj_p.paragraph_format.space_after = Pt(10)
    subj_run = subj_p.add_run(f"Re: Application for {job['title']}")
    _set_font(subj_run, size=11, bold=True, color=(0, 70, 127))

    # ── Salutation ────────────────────────────────────────────
    sal_p = doc.add_paragraph()
    sal_p.paragraph_format.space_after = Pt(8)
    sal_run = sal_p.add_run("Dear Hiring Manager,")
    _set_font(sal_run, size=11)

    # ── Company research note (subtle, woven into opening) ────
    company_note = cl_data.get("company_research_note", "")

    # ── Opening paragraph ─────────────────────────────────────
    open_p = doc.add_paragraph()
    open_p.paragraph_format.space_after = Pt(10)
    open_text = cl_data["opening_paragraph"]
    if company_note:
        open_text = open_text + " " + company_note
    open_run = open_p.add_run(open_text)
    _set_font(open_run, size=11)

    # ── Why me paragraph ─────────────────────────────────────
    why_p = doc.add_paragraph()
    why_p.paragraph_format.space_after = Pt(10)
    why_run = why_p.add_run(cl_data["why_me_paragraph"])
    _set_font(why_run, size=11)

    # ── 4 skills bullet points ─────────────────────────────────
    skills_intro_p = doc.add_paragraph()
    skills_intro_p.paragraph_format.space_after = Pt(4)
    intro_run = skills_intro_p.add_run(
        "In particular, I believe I can bring the following to your team:"
    )
    _set_font(intro_run, size=11)

    for item in cl_data.get("bullet_skills", []):
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.left_indent = Inches(0.4)
        bp.paragraph_format.space_after = Pt(4)

        skill_run = bp.add_run(f"{item['skill']}: ")
        _set_font(skill_run, size=11, bold=True)

        example_run = bp.add_run(item["example"])
        _set_font(example_run, size=11)

    doc.add_paragraph()  # spacer

    # ── Closing paragraph ─────────────────────────────────────
    close_p = doc.add_paragraph()
    close_p.paragraph_format.space_after = Pt(16)
    close_run = close_p.add_run(cl_data["closing_paragraph"])
    _set_font(close_run, size=11)

    # ── Sign-off ───────────────────────────────────────────────
    for line in ["Yours sincerely,", "", "", MASTER_CV["name"]]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        _set_font(run, size=11, bold=(line == MASTER_CV["name"]))

    doc.save(output_path)
    log.info(f"[OK] Cover letter saved: {output_path}")


# ═══════════════════════════════════════════════════════════════════
#  MAIN FUNCTION — called from job_search.py
# ═══════════════════════════════════════════════════════════════════

def generate_application_pack(job: dict, output_dir: str = ".") -> dict | None:
    """
    Generate a tailored CV + cover letter for a given job.
    Returns dict with file paths, or None if generation failed.
    Logs every step explicitly so failures are visible in GitHub Actions.
    """
    title   = job.get("title", "Unknown")
    company = job.get("company", "Unknown")
    log.info(f"[PACK] Starting: {title} @ {company}")

    # Step 1: Check dependencies
    log.info(f"[PACK] Step 1/4: Checking dependencies...")
    if not _check_dependencies():
        log.error(f"[PACK] FAILED Step 1 — missing packages. Check requirements.txt.")
        return None
    log.info(f"[PACK] Step 1/4: Dependencies OK")

    # Step 2: Check API key
    log.info(f"[PACK] Step 2/4: Checking API key...")
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        log.error(f"[PACK] FAILED Step 2 — ANTHROPIC_API_KEY not set in environment/secrets")
        return None
    log.info(f"[PACK] Step 2/4: API key present (length={len(api_key)})")

    # Step 3: Create output paths
    os.makedirs(output_dir, exist_ok=True)
    safe_name = re.sub(r"[^\w\s-]", "", f"{title}_{company}")
    safe_name = re.sub(r"\s+", "_", safe_name)[:50]
    cv_path = os.path.join(output_dir, f"CV_Ishaan_Kumar_{safe_name}.docx")
    cl_path = os.path.join(output_dir, f"CoverLetter_Ishaan_Kumar_{safe_name}.docx")
    log.info(f"[PACK] Step 3/4: Output paths: {cv_path} | {cl_path}")

    # Step 4: Call Claude API for tailored content
    log.info(f"[PACK] Step 4/4: Calling Claude API...")
    try:
        tailored = get_tailored_content(job)
    except Exception as e:
        log.error(f"[PACK] FAILED Step 4 — API call crashed: {e}")
        import traceback
        log.error(traceback.format_exc())
        return None

    if not tailored:
        log.error(f"[PACK] FAILED Step 4 — API returned no content for {title}")
        return None
    log.info(f"[PACK] Step 4/4: API returned tailored content OK")

    # Step 5: Build Word documents
    log.info(f"[PACK] Building CV docx...")
    try:
        build_cv_docx(job, tailored, cv_path)
        log.info(f"[PACK] CV saved: {cv_path} ({os.path.getsize(cv_path)} bytes)")
    except Exception as e:
        log.error(f"[PACK] CV build FAILED: {e}")
        import traceback
        log.error(traceback.format_exc())
        return None

    log.info(f"[PACK] Building cover letter docx...")
    try:
        build_cover_letter_docx(job, tailored, cl_path)
        log.info(f"[PACK] Cover letter saved: {cl_path} ({os.path.getsize(cl_path)} bytes)")
    except Exception as e:
        log.error(f"[PACK] Cover letter build FAILED: {e}")
        import traceback
        log.error(traceback.format_exc())
        return None

    log.info(f"[PACK] COMPLETE: {title} @ {company}")
    return {
        "job_title":    title,
        "company":      company,
        "cv_path":      cv_path,
        "cl_path":      cl_path,
        "ats_keywords": tailored.get("cv", {}).get("ats_keywords", []),
    }
